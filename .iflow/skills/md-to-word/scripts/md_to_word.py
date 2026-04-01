#!/usr/bin/env python3
"""
MD 转 WORD 转换器
将 Markdown 文件转换为 Word 文档（.docx），保留格式

用法：
    python md_to_word.py input.md
    python md_to_word.py input.md --output output.docx
    python md_to_word.py *.md  # 批量转换
"""

import argparse
import logging
import os
import re
import sys
from pathlib import Path
from typing import List, Optional


def parse_markdown_to_elements(md_content: str) -> List[dict]:
    """
    将 Markdown 内容解析为结构化元素列表

    Args:
        md_content: Markdown 文本内容

    Returns:
        元素列表，每个元素包含类型和内容
    """
    elements = []
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 标题（# H1, ## H2, ### H3, 等）
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line[level:].strip()
            elements.append({
                'type': 'heading',
                'level': level,
                'text': text
            })
            i += 1
            continue

        # 代码块（```）
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip() or 'text'
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append({
                'type': 'code',
                'lang': lang,
                'content': '\n'.join(code_lines)
            })
            i += 1
            continue

        # 水平线（--- 或 ***）
        if line.strip() in ['---', '***', '___']:
            elements.append({'type': 'hr'})
            i += 1
            continue

        # 引用（>）
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            elements.append({
                'type': 'quote',
                'text': quote_text
            })
            i += 1
            continue

        # 无序列表（- 或 *）
        if re.match(r'^\s*[-*]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*[-*]\s+', lines[i]):
                item_text = re.sub(r'^\s*[-*]\s+', '', lines[i]).strip()
                items.append(item_text)
                i += 1
            elements.append({
                'type': 'unordered_list',
                'items': items
            })
            continue

        # 有序列表（1.、2. 等）
        if re.match(r'^\s*\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                item_text = re.sub(r'^\s*\d+\.\s+', '', lines[i]).strip()
                items.append(item_text)
                i += 1
            elements.append({
                'type': 'ordered_list',
                'items': items
            })
            continue

        # 表格（Markdown 表格以 | 开头）
        if line.strip().startswith('|') and '|' in line.strip():
            # 解析表格
            rows = []
            # 表头
            header_row = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            rows.append(header_row)
            i += 1

            # 分隔行（包含 ---）
            if i < len(lines) and re.match(r'^\s*\|[\s\|:-]+\|\s*$', lines[i]):
                i += 1

            # 表格内容行
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [cell.strip() for cell in lines[i].strip().split('|')[1:-1]]
                if cells:  # 避免空行
                    rows.append(cells)
                i += 1

            if len(rows) > 1:  # 至少有表头和一行内容
                elements.append({
                    'type': 'table',
                    'rows': rows
                })
            continue

        # 普通段落
        paragraph_lines = []
        while i < len(lines) and lines[i].strip():
            # 检查是否是其他特殊元素
            if (lines[i].startswith('#') or
                lines[i].strip().startswith('```') or
                lines[i].strip() in ['---', '***', '___'] or
                lines[i].strip().startswith('>') or
                lines[i].strip().startswith('|') or
                re.match(r'^\s*[-*]\s+', lines[i]) or
                re.match(r'^\s*\d+\.\s+', lines[i])):
                break
            paragraph_lines.append(lines[i])
            i += 1

        if paragraph_lines:
            paragraph_text = ' '.join(paragraph_lines).strip()
            # 处理行内格式（粗体、斜体、代码、链接）
            paragraph_text = process_inline_formatting(paragraph_text)
            elements.append({
                'type': 'paragraph',
                'text': paragraph_text
            })

    return elements


def process_inline_formatting(text: str) -> str:
    """
    处理行内格式（粗体、斜体、代码、链接）

    Args:
        text: 原始文本

    Returns:
        处理后的文本（保留 Markdown 格式标记，Word 会处理）
    """
    # Word 会保留这些标记，后续在 Word 中可以进一步处理
    # 如果需要完全转换，可以使用更复杂的逻辑
    return text


def setup_logging() -> logging.Logger:
    """
    设置日志记录器，日志文件保存在 Python 脚本所在目录

    Returns:
        配置好的日志记录器
    """
    # 获取脚本所在目录
    script_dir = Path(__file__).parent
    log_path = script_dir / 'md_to_word.log'

    # 创建日志记录器
    logger = logging.getLogger('md_to_word')
    logger.setLevel(logging.DEBUG)

    # 清除已有的处理器
    logger.handlers.clear()

    # 文件处理器
    file_handler = logging.FileHandler(log_path, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    file_handler.setFormatter(file_formatter)

    # 控制台处理器
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_formatter = logging.Formatter('%(levelname)s: %(message)s')
    console_handler.setFormatter(console_formatter)

    # 添加处理器
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)

    logger.debug(f'日志文件: {log_path}')
    return logger


def convert_to_word(md_path: str, output_path: Optional[str] = None) -> str:
    """
    将 Markdown 文件转换为 Word 文档

    Args:
        md_path: Markdown 文件路径
        output_path: 输出 Word 文件路径（可选，默认使用原文件名）

    Returns:
        生成的 Word 文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        print("错误：缺少 python-docx 库")
        print("请安装：pip install python-docx")
        sys.exit(1)

    # 确定输出路径
    md_file = Path(md_path)
    if output_path is None:
        output_path = md_file.with_suffix('.docx')

    # 设置日志
    logger = setup_logging()
    logger.info(f'开始转换: {md_path}')
    logger.debug(f'输出路径: {output_path}')

    # 读取 Markdown 文件
    logger.debug(f'读取文件: {md_path}')
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    logger.debug(f'文件大小: {len(md_content)} 字符')

    # 解析 Markdown
    logger.debug('解析 Markdown 内容')
    elements = parse_markdown_to_elements(md_content)
    logger.debug(f'解析完成，共 {len(elements)} 个元素')

    # 创建 Word 文档
    logger.debug('创建 Word 文档')
    doc = Document()

    # 添加元素到 Word 文档
    element_counts = {
        'heading': 0,
        'paragraph': 0,
        'unordered_list': 0,
        'ordered_list': 0,
        'code': 0,
        'quote': 0,
        'table': 0,
        'hr': 0
    }

    for element in elements:
        if element['type'] == 'heading':
            # 标题
            level = min(element['level'], 3)  # Word 只支持 H1-H3
            heading = doc.add_heading(element['text'], level=level)
            heading.style.font.size = Pt(16 - level * 2)
            heading.style.font.bold = True
            element_counts['heading'] += 1

        elif element['type'] == 'paragraph':
            # 普通段落
            p = doc.add_paragraph(element['text'])
            p.style.font.size = Pt(11)
            element_counts['paragraph'] += 1

        elif element['type'] == 'unordered_list':
            # 无序列表
            for item in element['items']:
                p = doc.add_paragraph(item, style='List Bullet')
                p.style.font.size = Pt(11)
            element_counts['unordered_list'] += 1

        elif element['type'] == 'ordered_list':
            # 有序列表
            for item in element['items']:
                p = doc.add_paragraph(item, style='List Number')
                p.style.font.size = Pt(11)
            element_counts['ordered_list'] += 1

        elif element['type'] == 'code':
            # 代码块
            p = doc.add_paragraph()
            run = p.add_run(f"[代码: {element['lang']}]\n{element['content']}")
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.space_after = Pt(6)
            element_counts['code'] += 1

        elif element['type'] == 'quote':
            # 引用
            p = doc.add_paragraph(f"引用: {element['text']}")
            p.style.font.size = Pt(11)
            p.style.font.italic = True
            p.paragraph_format.left_indent = Pt(20)
            element_counts['quote'] += 1

        elif element['type'] == 'table':
            # 表格
            rows = element['rows']
            if not rows:
                continue

            # 创建表格
            num_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'

            # 填充表格内容
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    # 处理行内格式
                    cell_text = process_inline_formatting(cell_text)

                    # 如果是表头，加粗
                    if row_idx == 0:
                        table.rows[row_idx].cells[col_idx].text = cell_text
                        for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    else:
                        table.rows[row_idx].cells[col_idx].text = cell_text
            element_counts['table'] += 1

        elif element['type'] == 'hr':
            # 水平线
            p = doc.add_paragraph('_' * 80)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_after = Pt(12)
            element_counts['hr'] += 1

    logger.debug(f'元素统计: {element_counts}')

    # 保存 Word 文档
    logger.debug(f'保存 Word 文档: {output_path}')
    doc.save(output_path)
    logger.info(f'✓ 已转换: {md_path} -> {output_path}')
    logger.debug('转换完成')

    return str(output_path)


def convert_to_word_with_error_handling(md_path: str, output_path: Optional[str] = None) -> str:
    """
    带异常处理的转换函数

    Args:
        md_path: Markdown 文件路径
        output_path: 输出 Word 文件路径

    Returns:
        生成的 Word 文件路径

    Raises:
        Exception: 转换失败时抛出异常
    """
    try:
        return convert_to_word(md_path, output_path)
    except FileNotFoundError as e:
        logger = logging.getLogger('md_to_word')
        logger.error(f'文件未找到: {md_path} - {e}')
        raise
    except PermissionError as e:
        logger = logging.getLogger('md_to_word')
        logger.error(f'权限错误: {md_path} - {e}')
        raise
    except Exception as e:
        logger = logging.getLogger('md_to_word')
        logger.error(f'转换失败: {md_path} - {type(e).__name__}: {e}', exc_info=True)
        raise


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='将 Markdown 文件转换为 Word 文档',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python3 md_to_word.py README.md
  python3 md_to_word.py README.md --output README_word.docx
  python3 md_to_word.py *.md              # 批量转换
  python3 md_to_word.py docs/*.md         # 转换指定目录
        """
    )

    parser.add_argument(
        'input_files',
        nargs='+',
        help='输入 Markdown 文件路径（支持通配符）'
    )

    parser.add_argument(
        '--output', '-o',
        help='输出 Word 文件路径（仅限单个文件时使用）'
    )

    parser.add_argument(
        '--recursive', '-r',
        action='store_true',
        help='递归处理子目录'
    )

    args = parser.parse_args()

    # 设置全局日志
    main_logger = logging.getLogger('md_to_word_main')
    main_logger.setLevel(logging.INFO)
    main_logger.handlers.clear()

    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_formatter = logging.Formatter('%(levelname)s: %(message)s')
    console_handler.setFormatter(console_formatter)
    main_logger.addHandler(console_handler)

    main_logger.info('MD 转 WORD 转换器启动')

    # 展开通配符
    input_files = []
    for pattern in args.input_files:
        if '*' in pattern or '?' in pattern:
            matched = list(Path('.').glob(pattern))
            if args.recursive:
                matched.extend(Path('.').rglob(pattern.split('/')[-1]))
            input_files.extend(matched)
        else:
            input_files.append(Path(pattern))

    # 过滤不存在的文件
    valid_files = [f for f in input_files if f.exists() and f.suffix.lower() == '.md']

    if not valid_files:
        main_logger.error("错误：没有找到有效的 Markdown 文件")
        sys.exit(1)

    main_logger.info(f"找到 {len(valid_files)} 个 Markdown 文件")

    # 转换文件
    success_count = 0
    for md_file in valid_files:
        try:
            output_path = args.output if len(valid_files) == 1 and args.output else None
            convert_to_word(str(md_file), output_path)
            success_count += 1
        except Exception as e:
            main_logger.error(f"✗ 转换失败: {md_file} - {e}")

    main_logger.info(f"完成！成功转换 {success_count}/{len(valid_files)} 个文件")


if __name__ == '__main__':
    main()